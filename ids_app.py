import streamlit as st
import joblib
import pandas as pd
import numpy as np
from datetime import datetime
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import os

# ==============================
# PAGE CONFIG + STYLE
# ==============================
st.set_page_config(page_title="IDS Dashboard", layout="wide")

st.markdown("""
    <style>
    .stApp {
        background: linear-gradient(135deg, #0f2027, #203a43, #2c5364);
        color: white;
    }
    .css-1d391kg {background-color: #111;}
    .stButton>button {
        background-color: #00c6ff;
        color: black;
        font-weight: bold;
        border-radius: 10px;
    }
    </style>
""", unsafe_allow_html=True)

# ==============================
# LOAD MODEL
# ==============================
model = joblib.load("ids_model.pkl")
scaler = joblib.load("scaler.pkl")
features = joblib.load("features.pkl")

st.title("🚀 Intrusion Detection System (IDS)")
st.caption("Real-time Machine Learning-Based Threat Detection")

# ==============================
# SESSION STORAGE
# ==============================
if "results_df" not in st.session_state:
    st.session_state.results_df = pd.DataFrame()

# ==============================
# SIDEBAR
# ==============================
mode = st.sidebar.selectbox(
    "Select Mode",
    ["Manual Input", "Random Simulation", "Upload CSV"]
)

# ==============================
# PREDICTION FUNCTION
# ==============================
def predict(df):
    scaled = pd.DataFrame(scaler.transform(df), columns=features)
    preds = model.predict(scaled)
    probs = model.predict_proba(scaled)[:, 1]
    return preds, probs

# ==============================
# DISPLAY FUNCTION
# ==============================
def show_result(pred, prob):
    st.subheader("🔍 Prediction Result")

    col1, col2 = st.columns(2)

    with col1:
        st.metric("Attack Probability", f"{prob:.4f}")
        st.progress(float(prob))

    with col2:
        st.metric("Benign Probability", f"{1-prob:.4f}")
        st.progress(float(1-prob))

    if prob > 0.9:
        st.error("🚨 High Risk Attack")
    elif prob > 0.7:
        st.warning("⚠️ Medium Risk")
    else:
        st.success("✅ Benign")

# ==============================
# MANUAL INPUT
# ==============================
if mode == "Manual Input":

    st.subheader("🧮 Enter Features")

    input_data = {}
    cols = st.columns(3)

    for i, f in enumerate(features):
        with cols[i % 3]:
            input_data[f] = st.number_input(f, value=10.0)

    if st.button("Detect"):
        df = pd.DataFrame([input_data])
        pred, prob = predict(df)

        show_result(pred[0], prob[0])

        result = df.copy()
        result["Prediction"] = ["Attack" if pred[0] == 1 else "Benign"]
        result["Probability"] = prob

        st.session_state.results_df = pd.concat(
            [st.session_state.results_df, result],
            ignore_index=True
        )

# ==============================
# RANDOM SIMULATION
# ==============================
elif mode == "Random Simulation":

    st.subheader("🎲 Simulated Traffic")

    if st.button("Generate"):
        df = pd.DataFrame(
            [np.random.normal(50, 20, len(features))],
            columns=features
        )

        pred, prob = predict(df)

        st.write(df)
        show_result(pred[0], prob[0])

        df["Prediction"] = ["Attack" if pred[0] == 1 else "Benign"]
        df["Probability"] = prob

        st.session_state.results_df = pd.concat(
            [st.session_state.results_df, df],
            ignore_index=True
        )

# ==============================
# CSV UPLOAD
# ==============================
elif mode == "Upload CSV":

    file = st.file_uploader("Upload CSV", type=["csv"])

    if file:
        df = pd.read_csv(file)

        try:
            df = df[features]
            preds, probs = predict(df)

            df["Prediction"] = ["Attack" if p == 1 else "Benign" for p in preds]
            df["Probability"] = probs

            st.write(df.head())

            st.session_state.results_df = df

            st.bar_chart(df["Prediction"].value_counts())

        except Exception as e:
            st.error(e)

# ==============================
# RESULTS DISPLAY
# ==============================
st.subheader("📊 Results")

if not st.session_state.results_df.empty:
    st.dataframe(st.session_state.results_df)

# ==============================
# EXPORT TO WORD
# ==============================
def export_to_word(df):
    doc = Document()
    doc.add_heading('IDS Detection Report', 0)

    doc.add_paragraph(f"Generated: {datetime.now()}")

    # Table
    table = doc.add_table(rows=1, cols=len(df.columns))
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = col

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

    # Chart
    chart_path = "chart.png"
    df["Prediction"].value_counts().plot(kind='bar')
    plt.savefig(chart_path)
    plt.close()

    doc.add_picture(chart_path, width=Inches(5))

    file_path = "IDS_Report.docx"
    doc.save(file_path)

    return file_path

# ==============================
# DOWNLOAD BUTTON
# ==============================
if not st.session_state.results_df.empty:
    if st.button("📄 Export Report to Word"):
        file_path = export_to_word(st.session_state.results_df)

        with open(file_path, "rb") as f:
            st.download_button(
                "⬇️ Download Report",
                f,
                file_name="IDS_Report.docx"
            )
